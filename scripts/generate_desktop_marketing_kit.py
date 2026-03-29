import json
import os
import subprocess
import sys
import time
import uuid
from pathlib import Path
from urllib.error import URLError
from urllib.request import urlopen

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from playwright.sync_api import TimeoutError as PlaywrightTimeoutError
from playwright.sync_api import sync_playwright


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path(os.path.expanduser("~/Desktop"))
OUTPUT_DIR = DESKTOP / "智映视界官网与功能资料"
WEBSITE_URL = "https://www.zysj.site/"
LOCAL_BASE_URL = "http://127.0.0.1:5000"
APP_EXE = PROJECT_ROOT / "build" / "release" / "ZhiyingShijie" / "VideoFactory.exe"
WEBSITE_SCREENSHOT = DESKTOP / "智映视界官网截图.png"
DOC_PATH = DESKTOP / "智映视界功能软文推广.docx"


def wait_for_local_service(timeout_seconds: int = 30) -> None:
    deadline = time.time() + timeout_seconds
    last_error = ""
    while time.time() < deadline:
        try:
            with urlopen(f"{LOCAL_BASE_URL}/", timeout=5) as resp:
                if resp.status == 200:
                    return
        except URLError as exc:
            last_error = str(exc)
        except Exception as exc:
            last_error = str(exc)
        time.sleep(1)
    raise RuntimeError(f"本机绿色版服务未启动成功：{last_error}")


def ensure_local_app() -> None:
    try:
        with urlopen(f"{LOCAL_BASE_URL}/", timeout=3) as resp:
            if resp.status == 200:
                return
    except Exception:
        pass
    if not APP_EXE.exists():
        raise FileNotFoundError(f"未找到绿色版程序：{APP_EXE}")
    subprocess.Popen([str(APP_EXE)], cwd=str(APP_EXE.parent))
    wait_for_local_service()


def create_demo_account() -> tuple[str, str]:
    import requests

    username = f"codex_doc_{uuid.uuid4().hex[:8]}"
    password = "Codex123!"
    payload = {
        "username": username,
        "password": password,
        "accepted_agreements": True,
        "auto_login": True,
    }
    resp = requests.post(f"{LOCAL_BASE_URL}/api/auth/register", json=payload, timeout=20)
    resp.raise_for_status()
    data = resp.json()
    if not data.get("ok"):
        raise RuntimeError(data.get("error") or "注册演示账号失败")
    return username, password


def close_optional_modal(page, modal_id: str, close_selector: str) -> None:
    try:
        modal = page.locator(f"#{modal_id}")
        if modal.count() and modal.get_attribute("aria-hidden") == "false":
            page.locator(close_selector).click()
            page.wait_for_timeout(600)
    except Exception:
        return


def screenshot_locator(page, selector: str, output_path: Path) -> None:
    locator = page.locator(selector)
    locator.wait_for(state="visible", timeout=15000)
    locator.screenshot(path=str(output_path))


def show_panel(page, panel_id: str, section_id: str = "") -> None:
    page.evaluate(
        """([panelId, sectionId]) => {
            if (typeof showWorkspacePanel === 'function') {
                showWorkspacePanel(panelId);
            }
            if (sectionId && typeof activateHardSection === 'function') {
                activateHardSection(panelId, sectionId);
            }
        }""",
        [panel_id, section_id],
    )
    page.wait_for_timeout(1200)


def take_screenshots() -> dict[str, Path]:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    username, password = create_demo_account()

    screenshot_paths = {
        "website": WEBSITE_SCREENSHOT,
        "workspace": OUTPUT_DIR / "01_工作台总览.png",
        "assistant": OUTPUT_DIR / "02_智能助手.png",
        "mix_group": OUTPUT_DIR / "03_按组精准替换.png",
        "mix_mix": OUTPUT_DIR / "04_混剪裂变替换.png",
        "mix_partition": OUTPUT_DIR / "05_分区混剪裂变.png",
        "mix_sequence": OUTPUT_DIR / "06_槽位拼接混剪.png",
        "ai_make": OUTPUT_DIR / "07_AI智做.png",
        "effects": OUTPUT_DIR / "08_批量效果与资源库.png",
        "split": OUTPUT_DIR / "09_批量分割.png",
        "clip": OUTPUT_DIR / "10_片段微调.png",
        "manga": OUTPUT_DIR / "11_AI漫剧.png",
        "resource": OUTPUT_DIR / "12_资源互换.png",
        "account": OUTPUT_DIR / "13_账户中心与VIP.png",
    }

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        try:
            website_context = browser.new_context(viewport={"width": 1440, "height": 2200}, device_scale_factor=1.25)
            website_page = website_context.new_page()
            website_page.goto(WEBSITE_URL, wait_until="networkidle", timeout=30000)
            website_page.screenshot(path=str(screenshot_paths["website"]), full_page=True)
            website_context.close()

            user_context = browser.new_context(viewport={"width": 1600, "height": 1200}, device_scale_factor=1.25)
            page = user_context.new_page()
            page.goto(f"{LOCAL_BASE_URL}/user", wait_until="networkidle", timeout=30000)
            page.locator("#loginAccount").fill(username)
            page.locator("#loginPassword").fill(password)
            agreement = page.locator("#loginAgreementCheck")
            if not agreement.is_checked():
                agreement.check()
            page.locator("#loginForm button[type='submit']").click()
            page.wait_for_timeout(2000)
            close_optional_modal(page, "announcementModal", "[data-close-announcement='true']")
            page.wait_for_timeout(1200)

            screenshot_locator(page, "#workbenchApp", screenshot_paths["workspace"])
            show_panel(page, "panel-assistant")
            screenshot_locator(page, "#panel-assistant", screenshot_paths["assistant"])

            show_panel(page, "panel-materials")
            for strategy, key in (
                ("group", "mix_group"),
                ("mix", "mix_mix"),
                ("partition", "mix_partition"),
                ("sequence", "mix_sequence"),
            ):
                page.evaluate(
                    """(value) => {
                        if (typeof setMixStrategy === 'function') {
                            setMixStrategy(value);
                        }
                    }""",
                    strategy,
                )
                page.wait_for_timeout(900)
                screenshot_locator(page, "#panel-materials", screenshot_paths[key])

            show_panel(page, "panel-ai-make")
            screenshot_locator(page, "#panel-ai-make", screenshot_paths["ai_make"])

            show_panel(page, "panel-effects")
            screenshot_locator(page, "#panel-effects", screenshot_paths["effects"])

            show_panel(page, "panel-split")
            screenshot_locator(page, "#panel-split", screenshot_paths["split"])

            show_panel(page, "panel-clip")
            screenshot_locator(page, "#panel-clip", screenshot_paths["clip"])

            show_panel(page, "panel-ai-manga")
            screenshot_locator(page, "#panel-ai-manga", screenshot_paths["manga"])

            show_panel(page, "panel-resource-exchange", "resource-square-section")
            screenshot_locator(page, "#panel-resource-exchange", screenshot_paths["resource"])

            show_panel(page, "panel-account", "account-vip-section")
            screenshot_locator(page, "#panel-account", screenshot_paths["account"])
            user_context.close()
        finally:
            browser.close()

    return screenshot_paths


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(11)
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def add_title(doc: Document, text: str, size: int = 22) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(16 if level == 1 else 13)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(11)


def add_image(doc: Document, path: Path, width: float = 6.5) -> None:
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_word_doc(screenshots: dict[str, Path]) -> None:
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc, "智映视界功能推广稿")
    add_body(
        doc,
        "这是一套围绕剪映草稿批量生产的桌面工作台。它不是再造一套剪辑软件，"
        "而是把原本重复、分散、容易出错的换素材、换文字、配效果、整理草稿流程收在一个界面里，"
        "让内容团队更快把想法变成可直接继续编辑或交付的成片资产。"
    )

    add_heading(doc, "官网形象")
    add_body(
        doc,
        "官网首屏现在更强调“先试用核心功能，再决定让它变成你的视频生产线”。"
        "视觉上突出下载入口和能力概览，适合直接拿来做传播页、商务落地页和私域转化页。"
    )
    add_image(doc, screenshots["website"], width=6.2)

    sections = [
        (
            "工作台总览：从一个草稿延展成整条生产线",
            "真正有价值的不是把剪映搬进网页，而是把草稿、素材、导出、AI和账号体系放进一个统一工作台。"
            "用户进来之后只需要顺着左侧分类走，就能把“选草稿、批量替换、出新草稿、继续加工”这条主链路跑通。",
            screenshots["workspace"],
        ),
        (
            "智能助手：一句话先帮你找路，再帮你执行",
            "对新用户来说，最难的不是功能不够，而是不知道该先点哪里。命令中心把常见动作收成自然语言入口，"
            "可以先给出建议，再决定是否执行，降低学习成本，也更适合做演示和转化。",
            screenshots["assistant"],
        ),
        (
            "AI智做：把图生视频、文案、语音这些高频动作放进一条工作流",
            "AI 智做更像内容团队的中枢台，而不是单点工具集合。"
            "账号、模型、文案、图生视频和结果都被收在一个页面里，"
            "适合需要快速试创意、做图文转视频、做口播包装和补生成素材的场景。",
            screenshots["ai_make"],
        ),
        (
            "按组精准替换：适合有明确槽位对应关系的标准化项目",
            "这套模式最适合商品图、口播位、案例位、封面位都已经相对固定的项目。"
            "它会按草稿识别出来的槽位顺序严格对应素材目录，生成时每个槽位只取一个素材，"
            "所以成片结构稳定、出错率低，特别适合企业宣传、招商视频和标准化交付。",
            screenshots["mix_group"],
        ),
        (
            "混剪裂变替换：一个素材池，快速裂变出更多版本",
            "如果你的目标是做矩阵、做批量测试、做同结构多版本，这个模式会更直接。"
            "用户只要准备一个统一素材池，系统就能围绕同一参考草稿做不同组合，"
            "把原来需要反复试素材、反复复制工程的工作压成一次批量生成。",
            screenshots["mix_mix"],
        ),
        (
            "分区混剪裂变：片头、主体、片尾各自控制，不再互相串位",
            "很多内容不是简单随机混剪，而是需要片头、主体、片尾各自有自己的素材边界。"
            "分区混剪裂变就是为这种需求准备的，它既保留了裂变效率，又保证了关键结构不会乱，"
            "更适合做本地生活、招商项目、门店案例和带明确节奏的成片。",
            screenshots["mix_partition"],
        ),
        (
            "槽位拼接混剪：单槽位内连续拼视频，更适合做节奏型内容",
            "当一个槽位不该只放一条视频，而是要在同一个位置里连续拼接多段镜头时，"
            "槽位拼接混剪就能发挥优势。它把“单槽位内多片段组合”这件事独立出来，"
            "更适合做混剪号、节奏号、带货号和需要加强镜头密度的短视频矩阵。",
            screenshots["mix_sequence"],
        ),
        (
            "批量效果与资源库：内含14000+种丰富特效资源",
            "内容团队经常卡在“有草稿没质感”。资源库、特效配置和 Duo 资源接进来之后，"
            "就可以在原草稿基础上快速补转场、滤镜、人物优化和风格强化，不再反复跨软件找素材。",
            screenshots["effects"],
        ),
        (
            "批量分割：把文件分割、草稿处理和批量查看合成一个入口",
            "很多团队真正耗时间的不是做视频，而是做前处理。批量分割把文件切分、草稿结构查看和批量检查收在一起，"
            "适合做素材整理、主视频分段、镜头复检和导出前检查，能显著减少人工排查时间。",
            screenshots["split"],
        ),
        (
            "片段微调：把出片前最后一轮节奏和画面修正收在一起",
            "片段微调适合处理那些不值得整条重做、但又会明显影响观感的小问题。"
            "无论是节奏变速、画面校正还是局部镜头处理，都可以在最后阶段集中完成，"
            "让团队在保证效率的同时，仍然能把成片质感再往上提一档。",
            screenshots["clip"],
        ),
        (
            "AI漫剧：把脚本、分镜、场景素材和剪映草稿一次性串起来",
            "AI 漫剧是这套系统里最容易拉开差异化的能力。它不是停在脚本文案层，"
            "而是继续往后把场景目录、分镜信息和剪映草稿一起准备出来，让团队拿到的不是一个概念，"
            "而是一份可以直接继续填充、替换、调整、推进成片的工程底稿。"
            "对于剧情号、故事号、IP 口播包装、知识剧情演绎和批量短剧试题材，这条链路能明显缩短从创意到落地的距离。",
            screenshots["manga"],
        ),
        (
            "资源互换：把合作、置换和项目发布从聊天记录里拎出来",
            "资源互换页更像一个轻量协作广场。它适合团队去寻找合作方、发布自己的资源需求、查看互换机会，"
            "把原来散落在微信、QQ群、私聊里的资源对接，收成更可管理的一条业务线。",
            screenshots["resource"],
        ),
        (
            "账户与VIP：把试用、会员、授权和规则放到明面上",
            "商用软件最怕用户不知道自己买到了什么。账户中心把试用次数、VIP时效、卡类型、授权规则、邀请奖励都放到了一个入口，"
            "更适合销售演示、客户答疑，也更利于后续做正式商业交付。",
            screenshots["account"],
        ),
    ]

    for title, text, image_path in sections:
        add_heading(doc, title)
        add_body(doc, text)
        add_image(doc, image_path)

    add_heading(doc, "适合推广的话术方向")
    add_body(
        doc,
        "如果你已经在用剪映做批量内容，智映视界最适合用来承接“重复但不能出错”的那部分工作。"
        "它不要求团队重学一套工具，而是在你已经熟悉的草稿工作流上，把批量替换、批量混剪、效果补充、AI辅助和账号管理统一起来。"
    )
    add_body(
        doc,
        "对于本地生活、带货、矩阵号、剧情号和剪辑工作室来说，这套工具的价值不是炫技，"
        "而是把原来需要多个人、多次重复点击才能完成的流程压缩成标准动作，让团队真正能把“做视频”升级成“做视频生产线”。"
    )

    doc.save(str(DOC_PATH))


def main() -> None:
    ensure_local_app()
    screenshots = take_screenshots()
    build_word_doc(screenshots)
    print(
        json.dumps(
            {
                "website_screenshot": str(WEBSITE_SCREENSHOT),
                "output_dir": str(OUTPUT_DIR),
                "docx": str(DOC_PATH),
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    try:
        main()
    except PlaywrightTimeoutError as exc:
        raise SystemExit(f"playwright timeout: {exc}")
